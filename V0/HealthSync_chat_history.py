from docx import Document
from docx.shared import Pt

def add_code_block(doc, code, language="python"):
    para = doc.add_paragraph()
    run = para.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    para.style = doc.styles['Normal']
    para.paragraph_format.left_indent = Pt(18)
    para.paragraph_format.space_after = Pt(6)

def main():
    doc = Document()
    doc.add_heading('HealthSync Chat History', 0)
    doc.add_paragraph('This document contains the chat history between the user and Cascade (AI assistant) regarding the HealthSync project.')
    doc.add_paragraph('Date exported: 2025-04-17')
    doc.add_paragraph('---')

    # Chat 1
    doc.add_heading('User:', level=1)
    doc.add_paragraph('there is no gui for this make them i cant see whats happening there is no logs in gui make it structued so that i can see whats happening')
    doc.add_heading('Assistant:', level=1)
    doc.add_paragraph('To give you a clear, structured, and interactive GUI for model training (with real-time logs, progress, and status), I will:')
    doc.add_paragraph('- Create a new GUI script (e.g., gui_train.py) using tkinter.')
    doc.add_paragraph('- Display:')
    doc.add_paragraph('  - Real-time logs from TensorFlow/Keras training (epochs, loss, accuracy, etc.).')
    doc.add_paragraph('  - Buttons to start/stop training.')
    doc.add_paragraph('  - Model summary and configuration.')
    doc.add_paragraph('  - Optionally, a progress bar for epochs.')

    # Chat 2
    doc.add_heading('User:', level=1)
    doc.add_paragraph('it is using cpu to train the model can it use gpu to train it')
    doc.add_heading('Assistant:', level=1)
    doc.add_paragraph('Your current TensorFlow installation is not using the GPU. The result of tf.config.list_physical_devices(\'GPU\') is [], which means TensorFlow does not detect any GPU on your system.')
    doc.add_paragraph('How to Enable GPU Training with TensorFlow:')
    doc.add_paragraph('1. Check your hardware (NVIDIA GPU required).')
    doc.add_paragraph('2. Install NVIDIA GPU drivers.')
    doc.add_paragraph('3. Install CUDA 11.8 and cuDNN 8.6.')
    doc.add_paragraph('4. Add CUDA and cuDNN to PATH.')
    doc.add_paragraph('5. Run pip install tensorflow (already done).')
    doc.add_paragraph('6. Verify with:')
    add_code_block(doc, 'import tensorflow as tf\nprint(tf.config.list_physical_devices(\'GPU\'))')

    # Chat 3
    doc.add_heading('User:', level=1)
    doc.add_paragraph('convert all these chat into a pdf')
    doc.add_heading('Assistant:', level=1)
    doc.add_paragraph('I cannot directly export this chat to PDF from within this environment. However, I can guide you step-by-step to create a PDF of our conversation using the tools available on your system. (Instructions for copy-paste, Markdown, and print-to-PDF provided.)')

    # Chat 4
    doc.add_heading('User:', level=1)
    doc.add_paragraph('can you do for docx than pdf')
    doc.add_heading('Assistant:', level=1)
    doc.add_paragraph('Yes, I can help you create a .docx (Word) document with the content of our chat! (Instructions and a Python script template provided.)')

    # Chat 5
    doc.add_heading('User:', level=1)
    doc.add_paragraph('HealthSync Chat History\nThis document contains the chat history between the user and Cascade (AI assistant) regarding the HealthSync project.\nDate exported: 2025-04-17\n---\nUser:\nthere is no gui for this make them i cant see whats happening there is no logs in gui make it structued so that i can see whats happening\nAssistant:\nTo give you a clear, structured, and interactive GUI for model training (with real-time logs, progress, and status), I will: ...\nimport tkinter as tk\nfrom tkinter import scrolledtext, messagebox\n... only this much is there in docx')
    doc.add_heading('Assistant:', level=1)
    doc.add_paragraph('That’s because the script only had example entries. Please copy the full content above into your Word document or into the Python script as needed, and regenerate the .docx.')

    doc.save('HealthSync_chat_history.docx')

if __name__ == "__main__":
    main()
